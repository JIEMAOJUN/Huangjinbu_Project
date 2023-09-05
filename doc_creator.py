import docx



def docx_creator():
    doc = docx.Document()
    doc.add_paragraph('摘要')
    doc.add_paragraph('，对黄金埠电厂1号机组进行了深度调峰工况下一次调频性能试验。试验在机组长期运行所采取的控制模式下，分别在深度调峰工况37%额定负荷（Pe）、50%额定负荷（Pe）进行，所有工况试验时均维持机组正常运行稳定。通过记录数据，对该机组的一次调频滞后时间、幅度调整指标、动态转速不等率等进行了计算分析，给出了试验结论，并对存在问题进行了分析，提出了下一步的整改意见。')

    # 添加标题
    doc.add_heading('黄金埠电厂', level=0)
    # 添加章节标题
    doc.add_heading('1 前言', level=1)
    doc.add_paragraph('国能黄金埠电厂2号机组设计容量为600MW燃煤发电机组。汽轮机经高中压缸通流改造后铭牌出力均由原600MW扩容为650MW，汽轮机为上海汽轮机厂生产的超临界、反动式、一次中间再热、单轴、三缸四排汽、双背压、凝汽式汽轮机，发电机采用水氢氢的冷却方式；锅炉系上海锅炉厂生产的SG1913/25.4-M966型超临界一次中间再热、单炉膛、平衡通风、露天布置、固态排渣、全钢构架锅炉。燃烧系统采用四角切园燃烧系统。机组的监视与控制主要由上海西屋公司提供的OVATION集散控制系统（DCS）来实现。')
    doc.add_paragraph('为了考察黄金埠电厂2号机组参与电网运行的一次调频性能，同时确保发电机组参与电网一次调频时的调节安全，按照相关国家、行业标准与电网调度的要求，对该机组进行了深度调峰工况一次调频试验。 ')

    doc.add_heading('2 实验依据', level=1)
    doc.add_paragraph('a)  《火力发电机组一次调频试验及性能验收导则》（GB/T 30370-2013） ')
    doc.add_paragraph('b)  《电网运行准则》（GB/T 31464-2015）')
    doc.add_paragraph('c)  《发电机组并网安全条件及评价》（GB/T 28566-2012）')
    doc.add_paragraph('d)  《防止电力生产重大事故的二十五项重点要求》（国家能源局2014）')
    doc.add_paragraph('e)  《江西电网发电机组一次调频调度管理暂行规定》')
    doc.add_paragraph('f)   黄金埠电厂2号机组DCS 和 DEH 厂家相关资料')
    doc.add_paragraph('g)  《并网电源一次调频技术规定及试验导则》（GB/T40595-2021）')

    doc.add_heading('3 实验内容', level=1)
    doc.add_heading('3.1 电网对', level=2)
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')


    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        row = table.rows[i]
        for j in range(3):
            row.cells[j].text = f'({i+1}, {j+1})'

    doc.save('estimasting report.docx')

docx_creator()
